import os
import yaml
import zipfile
import json
import jsonschema
import jinja2
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from argparse import ArgumentParser


# Helper function to create a file tree
def generate_file_tree(schema_data, as_text=True):
    file_paths = []

    # Collect all file paths
    def collect_file_paths(element):
        for key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed', 'canContain']:
            if key in element:
                if key != "canContain":
                    for item in element[key]:
                        file_paths.append(item['path'])
                else:
                    for subelement in element[key]:
                        collect_file_paths(subelement)

    for element in schema_data.get('elements', []):
        collect_file_paths(element)

    # Create file tree structure (basic indentation)
    file_tree = {}
    for path in file_paths:
        parts = path.split('/')
        current = file_tree
        total = len(parts)
        for i, part in enumerate(parts):
            if i+1 < total:
                current = current.setdefault(f"{part}/", {})
            else:
                current = current.setdefault(f"{part}", {})

    # Convert file tree to string representation
    def tree_to_string(tree, indent=0):
        result = ''
        for key, subtree in tree.items():
            result += '  ' * indent + f"- {key}\n"
            result += tree_to_string(subtree, indent + 1)
        return result

    if as_text:
        return tree_to_string(file_tree)
    return file_tree


# Validate the schema itself
def validate_schema(schema_data):
    required_keys = ['name', 'description', 'version', 'elements']
    for key in required_keys:
        if key not in schema_data:
            return False, f"Missing required field: {key}"

    return True, "Schema is valid."


# Helper function to handle description output for Markdown
def format_description_markdown(description):
    if isinstance(description, list):
        return "\n\n".join(description)  # Each string in the list on a new line
    return description


# Helper function to handle description output for DOCX
def format_description_docx(doc, description):
    result = []
    if isinstance(description, list):
        for line in description:
            line = line.replace("<br>", "\n")
            result.append(doc.add_paragraph(line))  # Each string in the list as a separate paragraph
    else:
        result.append(doc.add_paragraph(description))
    return result

import os
import requests
import yaml
import json
from urllib.parse import urlparse

def load_schema(schema_path):
    parsed_url = urlparse(schema_path)
    if parsed_url.scheme in ("http", "https"):
        response = requests.get(schema_path)
        response.raise_for_status()  # Raises an error for bad status codes
        if schema_path.endswith('.jsonschema'):
            return response.json()
        elif schema_path.endswith('.zipschema'):
            return yaml.safe_load(response.text)
    else:
        with open(schema_path, 'r') as schema_file:
            if schema_path.endswith('.jsonschema'):
                return json.load(schema_file)
            elif schema_path.endswith('.zipschema'):
                return yaml.safe_load(schema_file)

    raise ValueError(f"Unsupported schema file format: {schema_path}")

# Helper function to validate a file against a JSON schema
def validate_jsonschema_file(zip_file, json_schema_path, file_path):
    with zip_file.open(file_path) as json_file:
        json_data = json.load(json_file)
        json_schema = load_schema(json_schema_path)
        try:
            jsonschema.validate(instance=json_data, schema=json_schema)
            return True, f"{file_path} is valid according to the JSON schema."
        except jsonschema.exceptions.ValidationError as e:
            return False, f"Validation failed for {file_path}: {str(e)}"

# Helper function to recursively validate using zipschema
def validate_zipschema(zip_file, zipschema_path, file_path):
    zipschema = load_schema(zipschema_path)
    # Recursively validate the zip file using the zipschema
    result, message = validate_zip_against_schema(zip_file.filename, zipschema)
    if result:
        return True, f"{file_path} validated successfully using zipschema."
    else:
        return False, f"Validation failed for {file_path} using zipschema: {message}"



# Validate a directory and its contents for canContain evaluator
def validate_canContain(zip_file, canContain, zip_contents):
    for item in canContain:
        directory_path = item.get('path', '')
        # Check if directory exists in the zip
        if any(content.startswith(directory_path) for content in zip_contents):
            # Apply any internal evaluators
            for evaluator in ['allOf', 'anyOf', 'oneOf', 'noneOf']:
                if evaluator in item:
                    result, message = validate_evaluator(zip_file, item[evaluator], evaluator, zip_contents,
                                                         directory_path)
                    if not result:
                        return False, message
        else:
            # If the directory is required but doesn't exist, raise an error
            if item.get('required', False):
                return False, f"Directory {directory_path} is required but not found in the zip file."
    return True, "canContain validation passed."


# General evaluator function
def validate_evaluator(zip_file, elements, evaluator_type, zip_contents, directory_path=""):
    one_of_found = 0
    for element in elements:
        element_path = os.path.join(directory_path, element['path']) if directory_path else element['path']
        if element_path in zip_contents:
            # Handle schema validation
            if 'schema' in element:
                schema_type = element['schema'].split('.')[-1].lower()
                if schema_type == 'jsonschema':
                    result, message = validate_jsonschema_file(zip_file, element['schema'], element_path)
                    if not result:
                        return False, message
                elif schema_type == 'zipschema' or element['schema'] == 'self':
                    if element['schema'] == 'self':
                        result, message = validate_zip_against_schema(zip_file.filename, elements)
                    else:
                        result, message = validate_zipschema(zip_file, element['schema'], element_path)
                    if not result:
                        return False, message

            # Check oneOf constraint
            if evaluator_type == 'oneOf':
                one_of_found += 1

    if evaluator_type == 'oneOf' and one_of_found != 1:
        return False, f"Exactly one file from 'oneOf' should be present in section: {evaluator_type}"

    return True, f"{evaluator_type} validation passed."


# Validate a zip file against the schema
def validate_zip_against_schema(zip_path, schema_data):
    if not zipfile.is_zipfile(zip_path):
        return False, "Provided file is not a valid zip file."

    with zipfile.ZipFile(zip_path, 'r') as zfile:
        zip_contents = zfile.namelist()

        for element in schema_data.get('elements', []):
            for key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'canContain']:
                if key == 'canContain':
                    result, message = validate_canContain(zfile, element['canContain'], zip_contents)
                    if not result:
                        return False, message
                else:
                    result, message = validate_evaluator(zfile, element.get(key, []), key, zip_contents)
                    if not result:
                        return False, message

    return True, "Zip file contents are valid."


# Function to add borders to DOCX table
def add_borders_to_cells(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'start', 'bottom', 'end']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # Border size
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)


# Generate Markdown documentation with canContain evaluator
def generate_markdown_with_tree(schema_data, output_file):
    file_tree = generate_file_tree(schema_data)

    template = """
# {{ schema.name }}

**Version:** {{ schema.version }}

**Description:**

{{ schema.description }}

## File Tree
{{ file_tree }}

## File List
| Filename | Summary | Section |
|----------|---------|---------|
{% for element in schema.elements %}
{% for key, items in element.items() if key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed', 'canContain'] %}
{% for item in items %}
| `{{ item.path }}` | {{ item.summary or '' }} | {{ element.section_title }} |
{% endfor %}
{% endfor %}
{% endfor %}

## Section List
{% for element in schema.elements %}
### {{ element.section_title }}
{{ format_description_markdown(element.section_description) }}

{% for key, items in element.items() if key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed', 'canContain'] %}
#### {{ key }}
{% for item in items %}
- **`{{ item.path }}`**: {{ item.description }} 
  {% if key == 'canContain' %}
    - Directory can contain the following:
    {% for subkey, subitems in item.items() if subkey in ['allOf', 'anyOf', 'oneOf', 'noneOf'] %}
      - **{{ subkey }}**:
        {% for subitem in subitems %}
          - `{{ subitem.path }}`: {{ subitem.description }}
        {% endfor %}
    {% endfor %}
  {% endif %}
{% endfor %}
{% endfor %}
{% endfor %}
"""

    env = jinja2.Environment(loader=jinja2.BaseLoader, trim_blocks=True, lstrip_blocks=True)
    template = env.from_string(template)
    if "description" in schema_data:
        schema_data["description"] = format_description_markdown(schema_data["description"])
    markdown_output = template.render(schema=schema_data, file_tree=file_tree,
                                      format_description_markdown=format_description_markdown)

    with open(output_file, 'w') as md_file:
        md_file.write(markdown_output)

    return "Markdown documentation with file table, sections, conditionals, and file items generated."


def set_keep_with_next(paragraph):
    """
    This function sets the 'keep with next' property on a paragraph.
    It keeps the paragraph together with the next paragraph.
    """
    pPr = paragraph._element.get_or_add_pPr()  # Get the paragraph's properties
    keepNext = OxmlElement('w:keepNext')  # Create a 'keepNext' element
    pPr.append(keepNext)  # Append it to the paragraph properties

def set_keep_together(paragraph):
    """
    This function sets the 'keep lines together' property on a paragraph.
    It prevents the paragraph from being split across pages.
    """
    pPr = paragraph._element.get_or_add_pPr()  # Get the paragraph's properties
    keepLines = OxmlElement('w:keepLines')  # Create a 'keepLines' element
    pPr.append(keepLines)  # Append it to the paragraph properties

def set_level(paragraph, level):
    from docx.shared import Inches
    paragraph.paragraph_format.left_indent = Inches(0.25 * level)  # Increase indent for each level

def set_bullet_level(paragraph, level):
    set_level(paragraph, level)
    paragraph.style = 'ListBullet'

def add_header(doc, header_text):
    from docx.shared import Pt
    section = doc.sections[0]  # Access the first section of the document
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = header_text
    header_paragraph.style = 'Header'  # Use the built-in 'Header' style


# Function to add a footer with a page number
def add_footer_with_page_number(doc):
    from docx.shared import Pt
    section = doc.sections[0]  # Access the first section of the document
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]

    # Create a page number field in the footer
    run = footer_paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)

    footer_paragraph.alignment = 2  # Align to the right

def add_styles(doc):
    from docx.shared import Pt
    # Add a new style to the document
    styles = doc.styles
    codestyle = styles.add_style('CodeStyle', 2)  # 1 represents a paragraph style

    # Set font to monospaced (e.g., Courier New)
    codestyle.font.name = 'Courier New'
    codestyle.font.size = Pt(10.5)

    # Set background color (light gray)
    for rPr in codestyle.element.xpath('.//w:rPr'):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D3D3D3')  # Hex code for light gray
        rPr.append(shd)

# Generate DOCX documentation with canContain evaluator
def generate_docx_with_tree(schema_data, output_file):
    doc = docx.Document()

    add_styles(doc)
    add_header(doc,schema_data['name'])
    add_footer_with_page_number(doc)

    # Title and version
    doc.add_heading(schema_data['name'], 0)
    doc.add_paragraph(f"Version: {schema_data['version']}")

    # Description
    format_description_docx(doc, schema_data['description'])

    doc.add_page_break()
    # File Tree Section
    set_keep_with_next(doc.add_heading('File Tree', level=1))
    file_tree = generate_file_tree(schema_data, as_text=False)

    def tree_to_doctree(tree, indent=0):
        for key, subtree in tree.items():
            p = doc.add_paragraph(f"{key}")
            set_keep_with_next(p)
            set_bullet_level(p, indent)
            tree_to_doctree(subtree, indent + 1)

    tree_to_doctree(file_tree,1)

    # File List Section (Table)
    doc.add_heading('File List', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filename'
    hdr_cells[1].text = 'Summary'
    hdr_cells[2].text = 'Section'

    for element in schema_data.get('elements', []):
        for key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed', 'canContain']:
            if key in element:
                for item in element[key]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item['path']
                    row_cells[1].text = item.get('summary', '')
                    row_cells[2].text = element['section_title']

    # Add borders to table
    add_borders_to_cells(table)
    doc.add_page_break()

    def add_section_text(item, level=0):
        p = doc.add_paragraph()
        set_keep_together(p)
        if item.get("description") is not None:
            p.add_run(f"{item['path']}\n\n    ").bold = True
            p.add_run(item['description'].replace("<br>", "\n"))
        else:
            if item.get("summary") is not None:
                p.add_run(f"{item['path']}: ").bold = True
                p.add_run(f"{item['summary']}")
            else:
                p.add_run(f"'{item['path']}' ").bold = True

        if item.get("example"):
            p.add_run(f"\n\nexample: \n    ").italic = True
            p.add_run(item.get("example"), style="CodeStyle")
        set_level(p, level)
        return p

    # Section List
    for element in schema_data.get('elements', []):
        # Section Title
        p = doc.add_heading(element['section_title'], level=1)
        set_keep_with_next(p)
        pees = format_description_docx(doc, element['section_description'])
        for p in pees:
            set_keep_with_next(p)

        # Conditional sub-sections (allOf, anyOf, oneOf, noneOf, allowed, canContain)
        for key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed', 'canContain']:
            if key in element:
                p = doc.add_heading(f"{key}: ", level=2)
                set_keep_together(p)
                set_keep_with_next(p)
                for item in element[key]:
                    if key == 'canContain':
                        p = add_section_text(item)
                        set_keep_with_next(p)
                        for subkey in ['allOf', 'anyOf', 'oneOf', 'noneOf']:
                            if subkey in item:
                                p = doc.add_heading(f"{subkey}:", level=3)
                                set_keep_with_next(p)
                                set_level(p, level=1)
                                for subitem in item[subkey]:
                                    add_section_text(subitem, 1)
                    else:
                        add_section_text(item,1)

    doc.save(output_file)

    return "DOCX documentation with file table, sections, conditionals, and file items generated."


# Command-line interface logic
def main():
    parser = ArgumentParser(description="Schema validator and zip file validator with documentation generator.")
    parser.add_argument('mode', choices=['validate-schema', 'validate-zip', 'generate-doc'],
                        help="The mode to run the tool in.")
    parser.add_argument('schema_file', help="Path to the zipschema YAML file.")
    parser.add_argument('--zipfile', help="Path to the zip file (required for zip validation).")
    parser.add_argument('--output', help="Path for the documentation output (optional for doc generation).")
    parser.add_argument('--format', choices=['markdown', 'docx'], help="Documentation format (markdown or docx).")

    args = parser.parse_args()

    # Load the schema file
    schema_data = load_schema(args.schema_file)

    # Handle the different modes
    if args.mode == 'validate-schema':
        valid, message = validate_schema(schema_data)
        print(message)

    elif args.mode == 'validate-zip':
        if not args.zipfile:
            print("Zip file path is required for zip validation.")
            return
        valid, message = validate_zip_against_schema(args.zipfile, schema_data)
        print(message)

    elif args.mode == 'generate-doc':
        if not args.format:
            print("Format is required for documentation generation.")
            return

        # Determine the output file location if not provided
        if not args.output:
            schema_dir = os.path.dirname(args.schema_file)
            schema_base_name = os.path.splitext(os.path.basename(args.schema_file))[0]
            if args.format == 'markdown':
                args.output = os.path.join(schema_dir, f"{schema_base_name}.md")
            elif args.format == 'docx':
                args.output = os.path.join(schema_dir, f"{schema_base_name}.docx")

        # Generate the appropriate documentation format
        if args.format == 'markdown':
            message = generate_markdown_with_tree(schema_data, args.output)
        elif args.format == 'docx':
            message = generate_docx_with_tree(schema_data, args.output)
        print(message)


if __name__ == '__main__':
    main()
